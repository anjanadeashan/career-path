import logging
import io
from flask import Blueprint, render_template, session, request, jsonify, send_file
from app.controllers.auth_helper import login_required
from app.repositories.resume_repository import ResumeRepository
from app.services.job_role_service import calculate_role_matches, CATEGORIES
from app.services.claude_service import ClaudeService
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)
job_bp = Blueprint('job', __name__)
resume_repo = ResumeRepository()
claude_service = ClaudeService()

@job_bp.route('/jobs', methods=['GET'])
@login_required
def jobs():
    """Display standard job role matches based on the user's skill profile."""
    user_id = session['user_id']

    resume = resume_repo.get_latest_resume(user_id)
    user_skills = resume_repo.get_skills_by_user(user_id)

    role_matches = calculate_role_matches(user_skills, user_id=user_id) if user_skills else []

    return render_template(
        'jobs.html',
        role_matches=role_matches,
        categories=CATEGORIES,
        has_resume=(resume is not None),
        has_skills=bool(user_skills)
    )

@job_bp.route('/api/generate-cover-letter', methods=['POST'])
@login_required
def generate_cover_letter():
    """Generate an AI cover letter for a specific job."""
    # Use get_json(silent=True) to prevent automatic 400 errors if headers are missing
    data = request.get_json(silent=True)
    
    if not data:
        return jsonify({"success": False, "error": "Invalid request format. Expected JSON."}), 400
        
    job_title = data.get('job_title')
    company_name = data.get('company_name')
    job_description = data.get('job_description', 'N/A')
    
    result = claude_service.generate_cover_letter(
        user_id=session['user_id'],
        job_title=job_title,
        company_name=company_name,
        job_description=job_description
    )
    return jsonify(result)

@job_bp.route('/api/download-cover-letter-word', methods=['POST'])
@login_required
def download_cover_letter_word():
    """Generate and download a Word document of the cover letter."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({"success": False, "error": "python-docx is not installed. Please run 'pip install python-docx'"}), 500
        
    data = request.get_json(silent=True)
    if not data or 'cover_letter_text' not in data:
        return jsonify({"success": False, "error": "Missing cover letter text."}), 400
        
    text = data.get('cover_letter_text')
    job_title = data.get('job_title', '').strip() or 'Cover Letter'
    company_name = data.get('company_name', '').strip() or 'Company'
    
    # Create a new Word Document
    doc = Document()
    
    # Set premium page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        for side in ('top', 'left', 'bottom', 'right'):
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), '6366f1')
            pgBorders.append(border)
        sectPr.append(pgBorders)
        
    # Set premium base font style
    style = doc.styles['Normal']
    style.font.name = 'Cambria'  # A professional, premium serif font
    style.font.size = Pt(11.5)
    style.font.color.rgb = RGBColor(0x1e, 0x20, 0x3a)

    # Top header for the premium document
    header_para = doc.add_paragraph(f'{job_title}  •  {company_name}')
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.runs[0]
    header_run.bold = True
    header_run.font.size = Pt(12)
    header_run.font.color.rgb = RGBColor(0x34, 0x3a, 0x54)
    header_para.paragraph_format.space_after = Pt(10)
    
    # Format the document text
    paragraphs = [paragraph.strip() for paragraph in text.split('\n') if paragraph.strip()]
    for paragraph in paragraphs:
        p = doc.add_paragraph(paragraph)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.2
            
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return send_file(
        file_stream,
        as_attachment=True,
        download_name='Cover_Letter.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
