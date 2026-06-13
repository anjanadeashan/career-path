import io
import json
import logging
import re
from datetime import date
from flask import Blueprint, render_template, request, session, redirect, url_for, flash, jsonify, send_file
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.controllers.auth_helper import login_required
from app.services.claude_service import ClaudeService
from app.repositories.job_repository import JobRepository
from app.repositories.resume_repository import ResumeRepository

logger = logging.getLogger(__name__)
career_bp = Blueprint('career', __name__)

claude_service = ClaudeService()
job_repo = JobRepository()
resume_repo = ResumeRepository()


def _parse_advisor_meta(feedback: dict) -> dict:
    """Extract enriched meta fields from the hidden HTML comment embedded in feedback_text."""
    if not feedback:
        return feedback
    raw = feedback.get('feedback_text', '')
    match = re.search(r'<!--advisor_meta:(\{.*?\})-->', raw, re.DOTALL)
    if match:
        try:
            meta = json.loads(match.group(1))
            feedback['profile_score'] = meta.get('profile_score', 0)
            feedback['top_skill_gaps'] = meta.get('top_skill_gaps', [])
            feedback['career_aspiration'] = meta.get('career_aspiration', {})
            feedback['transition_type'] = meta.get('transition_type', '')
            feedback['feedback_text'] = raw[:match.start()].rstrip()
        except (json.JSONDecodeError, KeyError):
            feedback.setdefault('profile_score', 0)
            feedback.setdefault('top_skill_gaps', [])
            feedback.setdefault('career_aspiration', {})
            feedback.setdefault('transition_type', '')
    else:
        feedback.setdefault('profile_score', 0)
        feedback.setdefault('top_skill_gaps', [])
        feedback.setdefault('career_aspiration', {})
        feedback.setdefault('transition_type', '')
    return feedback


# ── Word document helpers ────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _remove_table_borders(table):
    """Make a table invisible (no borders)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_page_border(section, color_hex: str = '6366f1', size: int = 12, space: int = 24):
    """Add a subtle page border to the document section."""
    sectPr = section._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:space'), str(space))
        border.set(qn('w:color'), color_hex)
        pgBorders.append(border)
    sectPr.append(pgBorders)


def _add_hr(doc, color_hex: str = 'e2e8f0', thickness: str = '6'):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), thickness)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def _set_doc_font(doc, name: str = 'Calibri'):
    """Set the default body font for the document."""
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1e, 0x20, 0x3a)


def _build_cover_letter_docx(cover_letter_text: str, job_title: str, company_name: str) -> bytes:
    """Build a premium .docx cover letter and return the raw bytes."""
    doc = Document()

    # ── Page setup ──────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)
    section.top_margin    = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    _set_page_border(section, color_hex='6366f1', size=12, space=24)

    _set_doc_font(doc)

    # ── Top accent banner (indigo table) ────────────────────────────────────
    banner = doc.add_table(rows=2, cols=1)
    _remove_table_borders(banner)

    # Row 0 — thick indigo stripe
    top_cell = banner.cell(0, 0)
    _set_cell_bg(top_cell, '6366f1')
    top_cell.height = Cm(0.25)
    top_para = top_cell.paragraphs[0]
    top_para.paragraph_format.space_before = Pt(0)
    top_para.paragraph_format.space_after  = Pt(0)

    # Row 1 — dark navy title row
    title_cell = banner.cell(1, 0)
    _set_cell_bg(title_cell, '0f1128')
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(14)
    title_para.paragraph_format.space_after  = Pt(6)
    t_run = title_para.add_run('COVER  LETTER')
    t_run.font.name  = 'Calibri Light'
    t_run.font.size  = Pt(26)
    t_run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
    t_run.font.bold  = False
    t_run.font.letter_spacing = Pt(4)  # wide-spaced feel

    # Subtitle inside banner
    sub_para = title_cell.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_before = Pt(0)
    sub_para.paragraph_format.space_after  = Pt(14)
    s_run = sub_para.add_run(f'{job_title}  ·  {company_name}')
    s_run.font.name  = 'Calibri'
    s_run.font.size  = Pt(10)
    s_run.font.color.rgb = RGBColor(0xa5, 0xb4, 0xfc)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)   # breathing room

    # ── Letter body ─────────────────────────────────────────────────────────
    lines = cover_letter_text.strip().split('\n')
    first_content = True

    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Blank separator between paragraphs
            spacer = doc.add_paragraph('')
            spacer.paragraph_format.space_before = Pt(0)
            spacer.paragraph_format.space_after  = Pt(8)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(10)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.2

        # Detect structural lines for accent styling
        is_date       = re.match(r'^(\d{1,2}[\s/.-]\w+|\w+ \d{1,2},?\s*\d{4}|Today|Date)', stripped, re.I)
        is_salutation = re.match(r'^(Dear |To Whom)', stripped, re.I)
        is_closing    = re.match(r'^(Sincerely|Yours (truly|faithfully|sincerely)|Best regards|Kind regards|Warm regards|Respectfully)', stripped, re.I)
        is_list_item  = stripped.startswith(('-', '*', '•'))

        run = p.add_run(stripped)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1e, 0x20, 0x3a)

        if is_date and first_content:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)  # slate
            run.font.size = Pt(10)
            first_content = False
        elif is_salutation:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.font.bold  = True
            run.font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)  # deep indigo
            first_content = False
        elif is_closing:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run.font.color.rgb = RGBColor(0x1e, 0x1b, 0x4b)
            run.font.bold = True
        elif is_list_item:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            first_content = False

    # ── Footer rule ──────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    _add_hr(doc, color_hex='6366f1', thickness='4')

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(6)
    footer_p.paragraph_format.space_after  = Pt(0)
    f_run = footer_p.add_run(f'Generated by CareerAI  ·  {date.today().strftime("%B %d, %Y")}')
    f_run.font.name  = 'Calibri'
    f_run.font.size  = Pt(9)
    f_run.font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)
    f_run.font.italic = True

    # ── Serialize ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ── Routes ───────────────────────────────────────────────────────────────────

@career_bp.route('/advisor', methods=['GET'])
@login_required
def advisor():
    """Display the AI Career Advisor feedback and suggestions page."""
    user_id = session['user_id']

    resume = resume_repo.get_latest_resume(user_id)
    feedback = job_repo.get_latest_career_feedback(user_id)
    feedback = _parse_advisor_meta(feedback)

    return render_template(
        'advisor.html',
        feedback=feedback,
        has_resume=(resume is not None)
    )


@career_bp.route('/advisor/generate', methods=['POST'])
@login_required
def generate_advice():
    """Trigger the Claude AI advisor pipeline to analyze CV and update advice."""
    user_id = session['user_id']

    result = claude_service.get_career_advice(user_id)

    if result.get('success'):
        flash("AI Career Advisor updated feedback successfully!", "success")
        return jsonify(result)
    else:
        error_msg = result.get('error', 'Failed to generate career advice.')
        flash(error_msg, "danger")
        return jsonify({"success": False, "error": error_msg}), 400


@career_bp.route('/advisor/cover-letter', methods=['POST'])
@login_required
def generate_cover_letter():
    """Generate a tailored cover letter using the user's resume and target job details."""
    user_id = session['user_id']
    data = request.get_json() or {}

    job_title    = data.get('job_title', '').strip()
    company_name = data.get('company_name', '').strip()
    job_description = data.get('job_description', '').strip()

    if not job_title or not company_name or not job_description:
        return jsonify({"success": False, "error": "Please fill in all fields: Job Title, Company Name, and Job Description."}), 400

    result = claude_service.generate_cover_letter(user_id, job_title, company_name, job_description)

    if result.get('success'):
        return jsonify(result)
    else:
        return jsonify({"success": False, "error": result.get('error', 'Failed to generate cover letter.')}), 400


@career_bp.route('/advisor/cover-letter/download', methods=['POST'])
@login_required
def download_cover_letter():
    """Stream a premium .docx cover letter file to the browser."""
    data = request.get_json() or {}

    cover_letter = data.get('cover_letter', '').strip()
    job_title    = data.get('job_title', 'Position').strip()
    company_name = data.get('company_name', 'Company').strip()

    if not cover_letter:
        return jsonify({'success': False, 'error': 'No cover letter text to download.'}), 400

    try:
        docx_bytes = _build_cover_letter_docx(cover_letter, job_title, company_name)
    except Exception as e:
        logger.error(f"Error building cover letter docx: {e}")
        return jsonify({'success': False, 'error': 'Failed to build document.'}), 500

    safe_company = re.sub(r'[^\w\s-]', '', company_name).replace(' ', '_')
    safe_title   = re.sub(r'[^\w\s-]', '', job_title).replace(' ', '_')
    filename     = f"Cover_Letter_{safe_company}_{safe_title}.docx"

    return send_file(
        io.BytesIO(docx_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )
