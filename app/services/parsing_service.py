import io
import re
import logging
import pypdf
import docx

logger = logging.getLogger(__name__)

class ParsingService:
    """Service to handle document ingestion and text extraction from PDF and DOCX files."""

    def extract_text(self, file_stream, file_name: str) -> str:
        """
        Extract raw text from a PDF or DOCX file stream.
        
        Args:
            file_stream: A file-like object (e.g., BytesIO or Flask FileStorage).
            file_name (str): The name of the file (to determine type).
            
        Returns:
            str: Extracted clean text.
        """
        file_extension = file_name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self._extract_pdf(file_stream)
        elif file_extension in ['docx', 'doc']:
            return self._extract_docx(file_stream)
        else:
            raise ValueError(f"Unsupported file format: .{file_extension}. Only PDF and DOCX files are allowed.")

    def _extract_pdf(self, stream) -> str:
        """Extract text from a PDF file stream using pypdf."""
        try:
            # Wrap in BytesIO if it is not already a file-like byte stream
            if isinstance(stream, bytes):
                stream = io.BytesIO(stream)
            
            reader = pypdf.PdfReader(stream)
            text = []
            
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                else:
                    logger.warning(f"Could not extract text from page {page_num + 1}")
            
            extracted_text = "\n".join(text).strip()
            if not extracted_text:
                raise ValueError("The PDF file seems to be empty or contains scanned images with no extractable text.")
            return self._fix_spaced_pdf_text(extracted_text)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise e

    def _fix_spaced_pdf_text(self, text: str) -> str:
        """Fix PDFs where every character is separated by a space: 'P y t h o n' -> 'Python'."""
        tokens = text.split()
        if not tokens or len(tokens) < 10:
            return text

        single_char_ratio = sum(1 for t in tokens if len(t) == 1) / len(tokens)
        if single_char_ratio < 0.4:
            return text  # Normal text, skip

        # Use double-spaces as word boundaries before collapsing
        result = re.sub(r'  +', '\x00', text)
        # Collapse single-space between alphanumeric chars (character spacing)
        result = re.sub(r'(?<=[A-Za-z0-9]) (?=[A-Za-z0-9])', '', result)
        # Also collapse space before punctuation attached to a word
        result = re.sub(r'(?<=[A-Za-z0-9]) (?=[.,;:!?])', '', result)
        # Restore word boundaries
        result = result.replace('\x00', ' ')
        # Normalize leftover multiple spaces
        result = re.sub(r' {2,}', ' ', result)

        logger.info("Applied character-spacing fix to PDF text (single-char ratio was %.0f%%)", single_char_ratio * 100)
        return result.strip()

    def _extract_docx(self, stream) -> str:
        """Extract text from a DOCX file stream using python-docx."""
        try:
            if isinstance(stream, bytes):
                stream = io.BytesIO(stream)
                
            doc = docx.Document(stream)
            text = []
            
            # Read paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
                    
            # Read tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text.append(" | ".join(row_text))
            
            extracted_text = "\n".join(text).strip()
            if not extracted_text:
                raise ValueError("The DOCX file seems to be empty or contains no extractable text.")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise e
